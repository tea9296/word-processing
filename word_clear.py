#!pip install python-docx
from docx import Document
from docx.enum.text import WD_BREAK
import os


def remove_header_and_footer(docx_path, save_path="./docs/"):
    # get the file name from docx_path
    file_name = docx_path.split('/')[-1]
    print(file_name)

    doc = Document(docx_path)

    doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

        for paragraph in section.even_page_header.paragraphs:
            paragraph.clear()
        for paragraph in section.even_page_footer.paragraphs:
            paragraph.clear()
        for paragraph in section.header.paragraphs:
            paragraph.clear()
        for paragraph in section.footer.paragraphs:
            paragraph.clear()

    doc.save(save_path + file_name)


## get all docx files i

## get all docx files in the folder
# target_dir = 'docs/產業報告/'
# destinated_dir = 'docs/產業報告p/'
# dess = os.listdir(destinated_dir)
# for file in os.listdir(target_dir):
#     if file in dess:
#         continue
#     if file.endswith('.docx'):
#         remove_header_and_footer(target_dir + file, save_path=destinated_dir)
